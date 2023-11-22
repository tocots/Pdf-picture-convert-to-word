        # programmer :kargilik的猫头鹰
        # 2023.08.09 in ZhongGuanCun,Beijing,China
import os
from PIL import Image
from tqdm import tqdm
import sys 
import fitz
from concurrent.futures import ThreadPoolExecutor
from docx import Document
from docx.shared import Inches
                                                                                          
def process_pdf(pdf_path):
    pdf_name = os.path.basename(pdf_path)  
    pdf_name_without_ext = os.path.splitext(pdf_name)[0]

    output_folder = os.path.join(os.path.dirname(pdf_path), pdf_name_without_ext)  
    os.makedirs(output_folder, exist_ok=True)

    pdf_doc = fitz.open(pdf_path)     
    page_count = pdf_doc.page_count  
    page_range = range(page_count)   
 
    def save_page_as_image(page_num):
        page = pdf_doc[page_num]
        pix = page.get_pixmap(dpi=300)
        pix.save(os.path.join(output_folder, f'page_{page_num+1:04d}.png'))
#Create a new folder with the pdf file name, and save each page in the pdf file as a png image to the folder
    with ThreadPoolExecutor(max_workers=8) as executor: 
        futures = [executor.submit(save_page_as_image, page_num) for page_num in tqdm(page_range)]

    for future in futures:
        future.result()

    pdf_doc.close()
    return output_folder
#Call multithreading to speed up program execution
def insert_images_into_word(doc, folder):
    doc.add_heading(os.path.basename(folder), level=3) #Use the pdf file name as the third-level title in the word document, and then place the generated picture under the heading3
    images = sorted(os.listdir(folder))  # sort the images to ensure they're in the correct order
    for image in images:
        if image.endswith('.png'):
            image_path = os.path.join(folder, image)
            with Image.open(image_path) as img:
                width, height = img.size
            image_width = Inches(6)
            scaled_height = image_width * height / width
            doc.add_picture(image_path, width=image_width, height=scaled_height)

# Skip the first argument, which is the script itself
folders = []
for pdf_path in sys.argv[1:]:
    print(f"Processing {pdf_path}...")
    folders.append(process_pdf(pdf_path))

doc = Document()                    # Generate Word document
for folder in folders:
    print(f"Inserting images from {folder} into Word document...")
    insert_images_into_word(doc, folder)
output_path = os.path.join(os.path.dirname(pdf_path), 'output.docx')
doc.save(output_path)
print("Word document saved as output.docx")

# Open the Word document
os.startfile(output_path)

print("All done! Good Job!")
input("Press Enter to exit...")
